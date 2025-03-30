import os
import json
import re
import tempfile
from typing import List, Dict, Any, Union, Optional
import numpy as np
from PIL import Image
from datetime import datetime
import requests
from io import BytesIO

# Optional imports for enhanced functionality
try:
    import tensorflow as tf
    TENSORFLOW_AVAILABLE = True
except ImportError:
    TENSORFLOW_AVAILABLE = False

try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False


class FileProcessor:
    """
    Handles processing of various file types uploaded by users.
    Supports extracting text from documents, analyzing images, and other file processing tasks.
    """
    
    def __init__(self, storage_dir: str = "uploads"):
        """
        Initialize the FileProcessor.
        
        Args:
            storage_dir: Directory to store uploaded files
        """
        self.storage_dir = storage_dir
        
        # Create storage directory if it doesn't exist
        os.makedirs(storage_dir, exist_ok=True)
        
        # Initialize image model if TensorFlow is available
        self.image_model = None
        self.image_model_available = False
        
        if TENSORFLOW_AVAILABLE:
            try:
                # Load a pre-trained model for image feature extraction
                self.image_model = tf.keras.applications.ResNet50(
                    include_top=False,
                    weights='imagenet',
                    pooling='avg'
                )
                self.image_model_available = True
            except Exception as e:
                print(f"TensorFlow model loading failed: {e}")
    
    def save_file(self, file_data: bytes, filename: str) -> str:
        """
        Save an uploaded file to disk.
        
        Args:
            file_data: Binary data of the file
            filename: Name of the file
            
        Returns:
            Path where the file was saved
        """
        # Create a safe filename
        safe_filename = ''.join(c for c in filename if c.isalnum() or c in '._-')
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        unique_filename = f"{timestamp}_{safe_filename}"
        
        # Save the file
        file_path = os.path.join(self.storage_dir, unique_filename)
        with open(file_path, 'wb') as f:
            f.write(file_data)
        
        return file_path
    
    def get_file_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from a file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary of metadata
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Basic metadata
        file_stats = os.stat(file_path)
        filename = os.path.basename(file_path)
        file_ext = os.path.splitext(filename)[1].lower()
        
        metadata = {
            "filename": filename,
            "file_type": file_ext.lstrip('.'),
            "size_bytes": file_stats.st_size,
            "created_at": datetime.fromtimestamp(file_stats.st_ctime).isoformat(),
            "modified_at": datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
        }
        
        # Additional metadata based on file type
        if file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']:
            try:
                img_metadata = self._get_image_metadata(file_path)
                metadata.update(img_metadata)
            except Exception as e:
                print(f"Error getting image metadata: {e}")
        
        elif file_ext in ['.pdf']:
            try:
                pdf_metadata = self._get_pdf_metadata(file_path)
                metadata.update(pdf_metadata)
            except Exception as e:
                print(f"Error getting PDF metadata: {e}")
                
        elif file_ext in ['.docx', '.doc']:
            try:
                doc_metadata = self._get_document_metadata(file_path)
                metadata.update(doc_metadata)
            except Exception as e:
                print(f"Error getting document metadata: {e}")
                
        elif file_ext in ['.csv', '.xlsx', '.xls']:
            try:
                sheet_metadata = self._get_spreadsheet_metadata(file_path)
                metadata.update(sheet_metadata)
            except Exception as e:
                print(f"Error getting spreadsheet metadata: {e}")
        
        return metadata
    
    def _get_image_metadata(self, image_path: str) -> Dict[str, Any]:
        """
        Extract metadata from an image file.
        
        Args:
            image_path: Path to the image
            
        Returns:
            Dictionary of image metadata
        """
        try:
            with Image.open(image_path) as img:
                metadata = {
                    "width": img.width,
                    "height": img.height,
                    "format": img.format,
                    "mode": img.mode,
                    "aspect_ratio": round(img.width / img.height, 3) if img.height > 0 else 0,
                }
                
                # Extract EXIF data if available
                if hasattr(img, '_getexif') and img._getexif():
                    exif = img._getexif()
                    if exif:
                        # Common EXIF tags
                        exif_tags = {
                            271: 'camera_make',
                            272: 'camera_model',
                            306: 'date_time',
                            36867: 'date_time_original',
                            33432: 'copyright',
                            37377: 'shutter_speed',
                            37378: 'aperture',
                            37379: 'brightness',
                            37380: 'exposure_compensation',
                            37383: 'metering_mode',
                            37384: 'flash',
                            37385: 'focal_length',
                            37386: 'focal_length_in_35mm',
                            34850: 'exposure_program',
                            34855: 'iso_speed',
                            41986: 'exposure_mode',
                            41987: 'white_balance',
                            41988: 'digital_zoom_ratio',
                            41990: 'scene_capture_type',
                            41991: 'gain_control',
                            41992: 'contrast',
                            41993: 'saturation',
                            41994: 'sharpness',
                        }
                        exif_data = {}
                        for tag, name in exif_tags.items():
                            if tag in exif:
                                exif_data[name] = str(exif[tag])
                        
                        if exif_data:
                            metadata['exif'] = exif_data
                
                return metadata
                
        except Exception as e:
            print(f"Error processing image metadata: {e}")
            return {}
    
    def _get_pdf_metadata(self, pdf_path: str) -> Dict[str, Any]:
        """
        Extract metadata from a PDF file.
        
        Args:
            pdf_path: Path to the PDF
            
        Returns:
            Dictionary of PDF metadata
        """
        if not PYPDF2_AVAILABLE:
            return {"error": "PyPDF2 library not available for PDF processing"}
        
        try:
            with open(pdf_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                metadata = {
                    "page_count": len(pdf_reader.pages),
                    "encrypted": pdf_reader.is_encrypted,
                }
                
                # Extract document info if available
                if pdf_reader.metadata:
                    info = pdf_reader.metadata
                    doc_info = {}
                    
                    # Common PDF metadata fields
                    if info.get('/Title'):
                        doc_info['title'] = info.get('/Title')
                    if info.get('/Author'):
                        doc_info['author'] = info.get('/Author')
                    if info.get('/Subject'):
                        doc_info['subject'] = info.get('/Subject')
                    if info.get('/Keywords'):
                        doc_info['keywords'] = info.get('/Keywords')
                    if info.get('/Creator'):
                        doc_info['creator'] = info.get('/Creator')
                    if info.get('/Producer'):
                        doc_info['producer'] = info.get('/Producer')
                    if info.get('/CreationDate'):
                        doc_info['creation_date'] = info.get('/CreationDate')
                    if info.get('/ModDate'):
                        doc_info['modification_date'] = info.get('/ModDate')
                        
                    if doc_info:
                        metadata['document_info'] = doc_info
                
                return metadata
                
        except Exception as e:
            print(f"Error processing PDF metadata: {e}")
            return {}
    
    def _get_document_metadata(self, doc_path: str) -> Dict[str, Any]:
        """
        Extract metadata from a Word document.
        
        Args:
            doc_path: Path to the document
            
        Returns:
            Dictionary of document metadata
        """
        if not DOCX_AVAILABLE:
            return {"error": "python-docx library not available for DOCX processing"}
        
        try:
            # Only handle .docx (not .doc) as python-docx doesn't support old format
            if not doc_path.lower().endswith('.docx'):
                return {"error": "Only .docx format is supported"}
            
            doc = docx.Document(doc_path)
            
            # Extract basic metadata
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "section_count": len(doc.sections),
                "table_count": len(doc.tables),
            }
            
            # Try to extract core properties if available
            try:
                core_properties = doc.core_properties
                prop_data = {}
                
                if hasattr(core_properties, 'author') and core_properties.author:
                    prop_data['author'] = core_properties.author
                if hasattr(core_properties, 'category') and core_properties.category:
                    prop_data['category'] = core_properties.category
                if hasattr(core_properties, 'comments') and core_properties.comments:
                    prop_data['comments'] = core_properties.comments
                if hasattr(core_properties, 'content_status') and core_properties.content_status:
                    prop_data['content_status'] = core_properties.content_status
                if hasattr(core_properties, 'created') and core_properties.created:
                    prop_data['created'] = core_properties.created.isoformat()
                if hasattr(core_properties, 'identifier') and core_properties.identifier:
                    prop_data['identifier'] = core_properties.identifier
                if hasattr(core_properties, 'keywords') and core_properties.keywords:
                    prop_data['keywords'] = core_properties.keywords
                if hasattr(core_properties, 'language') and core_properties.language:
                    prop_data['language'] = core_properties.language
                if hasattr(core_properties, 'last_modified_by') and core_properties.last_modified_by:
                    prop_data['last_modified_by'] = core_properties.last_modified_by
                if hasattr(core_properties, 'last_printed') and core_properties.last_printed:
                    prop_data['last_printed'] = core_properties.last_printed.isoformat()
                if hasattr(core_properties, 'modified') and core_properties.modified:
                    prop_data['modified'] = core_properties.modified.isoformat()
                if hasattr(core_properties, 'revision') and core_properties.revision:
                    prop_data['revision'] = core_properties.revision
                if hasattr(core_properties, 'subject') and core_properties.subject:
                    prop_data['subject'] = core_properties.subject
                if hasattr(core_properties, 'title') and core_properties.title:
                    prop_data['title'] = core_properties.title
                if hasattr(core_properties, 'version') and core_properties.version:
                    prop_data['version'] = core_properties.version
                
                if prop_data:
                    metadata['core_properties'] = prop_data
            
            except Exception as e:
                print(f"Error extracting document core properties: {e}")
            
            return metadata
            
        except Exception as e:
            print(f"Error processing document metadata: {e}")
            return {}
    
    def _get_spreadsheet_metadata(self, sheet_path: str) -> Dict[str, Any]:
        """
        Extract metadata from a spreadsheet file.
        
        Args:
            sheet_path: Path to the spreadsheet
            
        Returns:
            Dictionary of spreadsheet metadata
        """
        if not PANDAS_AVAILABLE:
            return {"error": "pandas library not available for spreadsheet processing"}
        
        try:
            # Handle different formats
            file_ext = os.path.splitext(sheet_path)[1].lower()
            
            if file_ext == '.csv':
                df = pd.read_csv(sheet_path)
                metadata = {
                    "row_count": len(df),
                    "column_count": len(df.columns),
                    "column_names": list(df.columns),
                    "file_type": "csv"
                }
            elif file_ext in ['.xlsx', '.xls']:
                # Get sheet names without loading the whole file
                sheet_names = pd.ExcelFile(sheet_path).sheet_names
                
                # Load the first sheet for basic metrics
                df = pd.read_excel(sheet_path, sheet_name=0)
                
                metadata = {
                    "sheet_count": len(sheet_names),
                    "sheet_names": sheet_names,
                    "row_count_first_sheet": len(df),
                    "column_count_first_sheet": len(df.columns),
                    "column_names_first_sheet": list(df.columns),
                    "file_type": "excel"
                }
            else:
                return {"error": f"Unsupported spreadsheet format: {file_ext}"}
            
            return metadata
            
        except Exception as e:
            print(f"Error processing spreadsheet metadata: {e}")
            return {}
    
    def extract_text_from_file(self, file_path: str) -> str:
        """
        Extract text content from a file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text content
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # Text files
        if file_ext in ['.txt', '.md', '.csv', '.json', '.xml', '.html', '.htm']:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except UnicodeDecodeError:
                # Try with a different encoding if utf-8 fails
                try:
                    with open(file_path, 'r', encoding='latin-1') as f:
                        return f.read()
                except Exception as e:
                    print(f"Error reading text file: {e}")
                    return ""
        
        # PDF files
        elif file_ext == '.pdf':
            return self._extract_text_from_pdf(file_path)
        
        # Word documents
        elif file_ext in ['.docx', '.doc']:
            return self._extract_text_from_document(file_path)
        
        # Images (OCR)
        elif file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']:
            return self._extract_text_from_image(file_path)
        
        # Unsupported format
        else:
            print(f"Unsupported file format for text extraction: {file_ext}")
            return ""
    
    def _extract_text_from_pdf(self, pdf_path: str) -> str:
        """
        Extract text from a PDF file.
        
        Args:
            pdf_path: Path to the PDF
            
        Returns:
            Extracted text
        """
        if not PYPDF2_AVAILABLE:
            return "PDF text extraction not available (PyPDF2 library missing)"
        
        try:
            text = []
            with open(pdf_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                # Extract text from each page
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text.append(page.extract_text())
            
            return "\n".join(text)
            
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return ""
    
    def _extract_text_from_document(self, doc_path: str) -> str:
        """
        Extract text from a Word document.
        
        Args:
            doc_path: Path to the document
            
        Returns:
            Extracted text
        """
        if not DOCX_AVAILABLE:
            return "Document text extraction not available (python-docx library missing)"
        
        try:
            # Only handle .docx (not .doc)
            if not doc_path.lower().endswith('.docx'):
                return "Only .docx format is supported for text extraction"
            
            doc = docx.Document(doc_path)
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    full_text.append(" | ".join(row_text))
            
            return "\n".join(full_text)
            
        except Exception as e:
            print(f"Error extracting text from document: {e}")
            return ""
    
    def _extract_text_from_image(self, image_path: str) -> str:
        """
        Extract text from an image using OCR.
        
        Args:
            image_path: Path to the image
            
        Returns:
            Extracted text
        """
        if not TESSERACT_AVAILABLE:
            return "Image text extraction not available (pytesseract library missing)"
        
        try:
            img = Image.open(image_path)
            
            # Convert to RGB if necessary (some formats like PNG with transparency can cause issues)
            if img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Apply some preprocessing to improve OCR results
            # 1. Resize for better OCR (if image is very large or small)
            width, height = img.size
            if width > 2000 or height > 2000:
                ratio = min(2000/width, 2000/height)
                new_width = int(width * ratio)
                new_height = int(height * ratio)
                img = img.resize((new_width, new_height), Image.LANCZOS)
            
            # 2. Use Tesseract to extract text
            text = pytesseract.image_to_string(img)
            return text
            
        except Exception as e:
            print(f"Error extracting text from image: {e}")
            return ""
    
    def extract_image_features(self, image_path: str) -> Optional[np.ndarray]:
        """
        Extract feature vector from an image using a pre-trained model.
        
        Args:
            image_path: Path to the image
            
        Returns:
            Feature vector as numpy array, or None if extraction fails
        """
        if not self.image_model_available:
            print("Image feature extraction not available (TensorFlow/model missing)")
            return None
        
        try:
            # Load and preprocess the image
            img = Image.open(image_path)
            
            # Convert to RGB if necessary
            if img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Resize to the expected input size for the model
            img = img.resize((224, 224))
            
            # Convert to numpy array and preprocess
            img_array = tf.keras.preprocessing.image.img_to_array(img)
            img_array = np.expand_dims(img_array, axis=0)
            img_array = tf.keras.applications.resnet50.preprocess_input(img_array)
            
            # Extract features
            features = self.image_model.predict(img_array)
            return features[0]  # Return the feature vector
            
        except Exception as e:
            print(f"Error extracting image features: {e}")
            return None
    
    def get_file_summary(self, file_path: str) -> Dict[str, Any]:
        """
        Generate a comprehensive summary of a file including metadata and content overview.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary with file summary information
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Get basic metadata
        metadata = self.get_file_metadata(file_path)
        
        # Get content summary
        file_ext = os.path.splitext(file_path)[1].lower()
        
        summary = {
            "metadata": metadata,
            "content_preview": None,
            "extracted_text_length": 0,
            "has_searchable_text": False,
        }
        
        # Extract text for content preview
        try:
            extracted_text = self.extract_text_from_file(file_path)
            if extracted_text:
                # Set flag for searchable text
                summary["has_searchable_text"] = True
                
                # Store the length of extracted text
                summary["extracted_text_length"] = len(extracted_text)
                
                # Create a preview (first 500 characters)
                preview_length = min(500, len(extracted_text))
                summary["content_preview"] = extracted_text[:preview_length]
                
                # Try to detect the main language
                try:
                    # Simple language detection based on common words
                    lang_markers = {
                        'the': 'en', 'and': 'en', 'of': 'en',  # English
                        'der': 'de', 'die': 'de', 'und': 'de',  # German
                        'le': 'fr', 'la': 'fr', 'les': 'fr',    # French
                        'el': 'es', 'los': 'es', 'las': 'es',   # Spanish
                        '的': 'zh', '是': 'zh', '在': 'zh',      # Chinese
                        'の': 'ja', 'に': 'ja', 'は': 'ja',      # Japanese
                    }
                    
                    words = re.findall(r'\b\w+\b', extracted_text.lower())
                    lang_counts = {}
                    
                    for word in words:
                        if word in lang_markers:
                            lang = lang_markers[word]
                            lang_counts[lang] = lang_counts.get(lang, 0) + 1
                    
                    if lang_counts:
                        main_language = max(lang_counts.items(), key=lambda x: x[1])[0]
                        summary["detected_language"] = main_language
                except Exception as e:
                    print(f"Language detection error: {e}")
        
        except Exception as e:
            print(f"Error creating content summary: {e}")
        
        return summary